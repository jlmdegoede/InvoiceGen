from docx import Document
from docx.shared import Inches
from Utils.date_helper import get_formatted_string
from docx.shared import Pt


def generate_docx_invoice(invoice, user, products, tax_rate):
    document = Document()
    p = document.add_paragraph()
    p.add_run('Adres: \t\t')
    p.add_run(user.name).bold = True
    p.add_run('\n\t\t' + user.address)
    p.add_run('\n\t\t' + user.city_and_zipcode)
    p.add_run('\nE-mail:\t\t' + user.email)
    p.add_run('\nIBAN:\t\t' + user.iban)
    if user.kvk:
        p.add_run('\nKvK:\t\t' + user.kvk)
    if user.btw_number:
        p.add_run('\nBTW-nr.: \t' + user.btw_number)

    p = document.add_paragraph()
    p.add_run(invoice.to_company.company_name)
    p.add_run('\n' + invoice.to_company.company_address)
    p.add_run('\n' + invoice.to_company.company_city_and_zipcode)

    document.add_heading(invoice.title, 0)
    p = document.add_paragraph()
    p.add_run('\nVolgnummer: \t\t' + str(invoice.invoice_number))
    p.add_run('\nFactuurdatum: \t' + get_formatted_string(invoice.date_created))
    p.add_run('\nVervaldatum: \t\t' + get_formatted_string(invoice.expiration_date))
    document.add_paragraph()

    table = document.add_table(rows=1, cols=5)
    hdr_cells = table.rows[0].cells
    hdr_cells[0].paragraphs[0].add_run('Opdracht').bold = True
    hdr_cells[1].paragraphs[0].add_run('Identificatienr.').bold = True
    hdr_cells[2].paragraphs[0].add_run('Kwantiteit').bold = True
    hdr_cells[3].paragraphs[0].add_run('Prijs per eenheid').bold = True
    hdr_cells[4].paragraphs[0].add_run('Prijs').bold = True
    for product in products:
        row_cells = table.add_row().cells
        row_cells[0].text = str(product.title)
        row_cells[1].text = str(product.identification_number)
        row_cells[2].text = str(product.quantity)
        row_cells[3].text = '€' + str(product.price_per_quantity)
        row_cells[4].text = '€' + str(product.get_price())

    document.add_paragraph()

    if tax_rate:
        p = document.add_paragraph()
        run = p.add_run()
        run.add_break()
        p.add_run('Subtotaal: \t €' + str(invoice.get_totaalbedrag()))
        p.alignment = 2
        p.add_run('\nBTW: \t\t €' + str(invoice.get_btw()))
        p.alignment = 2
        p.add_run('\nTotaal te voldoen: \t €' + str(invoice.get_totaalbedrag() + invoice.get_btw()))
        p.alignment = 2
    else:
        document.add_paragraph('Totaal te voldoen: \t €' + str(invoice.get_totaalbedrag()))

    document.add_paragraph()
    document.add_paragraph('Gelieve uw betaling uiterlijk te voldoen op ' + get_formatted_string(invoice.expiration_date) + ' op IBAN ' + user.iban)

    document.add_page_break()
    return document